import json
from io import StringIO

from flask import Flask, jsonify, request
from flask_cors import CORS
from textblob import TextBlob

app = Flask(__name__)
CORS(app)
answers = []


@app.route('/health')
def health():
    return jsonify({"status": "ok"})




def load_questions_short():
    questions_short = []
    with open('data/questions_short.txt', encoding="utf8") as f:
        questions_short = f.readlines()


    return questions_short


def getContractAnalysis(selected_response):
    print(selected_response)
    
    if selected_response == "":
        return "No answer found in document"
    else:
        blob = TextBlob(selected_response)
        polarity = blob.sentiment.polarity
        print(polarity)

        if polarity > 0:
            return "Positive"
        elif polarity < 0:
            return "Negative"
        else:
            return "Neutral"



questions_short = load_questions_short()



@app.route('/questionsshort')
def getQuestionsShort():
    # Ensure proper JSON response
    return jsonify(questions_short)




from agent import LegalAgent
import pypdf
import docx
import hashlib

agent = LegalAgent()
last_processed_hash = None

@app.route('/contracts/', methods=["POST"])
def getContractResponse():
    question = request.form.get('question', '')
    
    if 'file' in request.files and request.files['file'].filename != '':
        file = request.files["file"]
        
        paragraph = ""
        if file.filename.lower().endswith('.pdf'):
            try:
                reader = pypdf.PdfReader(file.stream)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        paragraph += text + "\n"
            except Exception as e:
                print(f"Error parsing PDF: {e}")
        elif file.filename.lower().endswith('.docx'):
            try:
                doc = docx.Document(file.stream)
                for p in doc.paragraphs:
                    if p.text:
                        paragraph += p.text + "\n"
            except Exception as e:
                print(f"Error parsing DOCX: {e}")
        else:
            paragraph = file.read().decode("utf-8", errors="replace")

        
        # Cache mechanism to prevent re-embedding the same document on every question
        if len(paragraph) > 0:
            doc_hash = hashlib.md5(paragraph.encode('utf-8')).hexdigest()
            global last_processed_hash
            if doc_hash != last_processed_hash:
                print("Embedding document into Vector DB...")
                metadata = agent.process_document(paragraph)
                last_processed_hash = doc_hash
            else:
                print("Document already embedded. Skipping FAISS generation...")
            
    if not question:
        return "Please ask a question."

    print("Running Agent...")
    result = agent.run(question)
    
    # We return the new structured result
    # We format it to match the UI expectations or update the UI to handle it.
    return jsonify(result)

@app.route('/contracts/upload', methods=["POST"])
def uploadContract():
    if 'file' in request.files and request.files['file'].filename != '':
        file = request.files["file"]
        paragraph = ""
        if file.filename.lower().endswith('.pdf'):
            try:
                reader = pypdf.PdfReader(file.stream)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        paragraph += text + "\n"
            except Exception as e:
                print(f"Error parsing PDF: {e}")
        elif file.filename.lower().endswith('.docx'):
            try:
                doc = docx.Document(file.stream)
                for p in doc.paragraphs:
                    if p.text:
                        paragraph += p.text + "\n"
            except Exception as e:
                print(f"Error parsing DOCX: {e}")
        else:
            paragraph = file.read().decode("utf-8", errors="replace")

        if len(paragraph) > 0:
            doc_hash = hashlib.md5(paragraph.encode('utf-8')).hexdigest()
            global last_processed_hash
            metadata = agent.process_document(paragraph)
            last_processed_hash = doc_hash
            return jsonify(metadata)
            
    return jsonify({"error": "No valid file uploaded"}), 400


if __name__ == '__main__':
    app.run(port=5001)